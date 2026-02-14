import io
from datetime import date
import streamlit as st

from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

from docx import Document
from docx.shared import Inches

st.set_page_config(page_title="MAKK Invoice Generator", layout="centered")

st.title("MAKK Invoice Generator")

# --- Inputs ---
col1, col2 = st.columns(2)
with col1:
    inv_date = st.date_input("Date", value=date.today())
    invoice_no = st.text_input("Invoice # (or Customer ID)", value="")
with col2:
    receiver = st.text_input("To (Receiver name / Company)", value="")
    phone = st.text_input("Phone", value="")

address = st.text_area("Address", value="")

st.subheader("Line Item")
c1, c2 = st.columns(2)
with c1:
    quantity = st.number_input("Quantity", min_value=0.0, step=1.0, value=1.0)
    weight = st.number_input("Weight", min_value=0.0, step=0.1, value=0.0)
with c2:
    description = st.text_input("Description", value="")
    line_total = st.number_input("Line Total (USD) - manual", min_value=0.0, step=1.0, value=0.0)

st.subheader("Totals")
t1, t2 = st.columns(2)
with t1:
    subtotal = st.number_input("Subtotal (USD) - manual", min_value=0.0, step=1.0, value=0.0)
with t2:
    sales_tax = st.number_input("Sales Tax (USD) - manual", min_value=0.0, step=1.0, value=0.0)

total = round(subtotal + sales_tax, 2)
st.write(f"**Total (auto): ${total:.2f}**")

# --- Helper: Generate PDF ---
def build_pdf():
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=LETTER)
    width, height = LETTER

    # Logo (optional)
    # Put your logo at assets/logo.png in your repo
    try:
        c.drawImage("assets/logo.png", 0.75*inch, height - 1.25*inch, width=1.2*inch, height=1.2*inch, mask="auto")
    except Exception:
        pass

    # Header text
    c.setFont("Helvetica-Bold", 14)
    c.drawString(2.2*inch, height - 0.9*inch, "MAKK CROSS BORDER SOLUTIONS LTD.")

    c.setFont("Helvetica", 10)
    c.drawString(2.2*inch, height - 1.1*inch, "address | 626-601-613")

    y = height - 1.6*inch
    c.setFont("Helvetica-Bold", 12)
    c.drawString(0.75*inch, y, "INVOICE")
    y -= 0.3*inch

    c.setFont("Helvetica", 10)
    c.drawString(0.75*inch, y, f"Date: {inv_date.isoformat()}")
    c.drawString(3.8*inch, y, f"Invoice #: {invoice_no}")
    y -= 0.25*inch

    c.drawString(0.75*inch, y, f"To: {receiver}")
    y -= 0.2*inch
    c.drawString(0.75*inch, y, f"Phone: {phone}")
    y -= 0.2*inch

    # Address multi-line
    c.drawString(0.75*inch, y, "Address:")
    y -= 0.2*inch
    for line in (address or "").splitlines():
        c.drawString(1.1*inch, y, line)
        y -= 0.18*inch

    y -= 0.15*inch
    c.setFont("Helvetica-Bold", 10)
    c.drawString(0.75*inch, y, "Qty")
    c.drawString(1.3*inch, y, "Description")
    c.drawString(4.6*inch, y, "Weight")
    c.drawString(5.5*inch, y, "Amount")
    y -= 0.2*inch

    c.setFont("Helvetica", 10)
    c.drawString(0.75*inch, y, f"{quantity:g}")
    c.drawString(1.3*inch, y, description[:60])
    c.drawString(4.6*inch, y, f"{weight:g}")
    c.drawString(5.5*inch, y, f"${line_total:,.2f}")
    y -= 0.4*inch

    # Totals
    c.drawString(4.6*inch, y, "Subtotal:")
    c.drawRightString(7.5*inch, y, f"${subtotal:,.2f}")
    y -= 0.2*inch
    c.drawString(4.6*inch, y, "Sales Tax:")
    c.drawRightString(7.5*inch, y, f"${sales_tax:,.2f}")
    y -= 0.2*inch

    c.setFont("Helvetica-Bold", 11)
    c.drawString(4.6*inch, y, "Total:")
    c.drawRightString(7.5*inch, y, f"${total:,.2f}")
    y -= 0.5*inch

    # Footer note
    c.setFont("Helvetica", 10)
    c.drawString(0.75*inch, y, "Note: Make all checks payable to MAKK CROSS BORDER SOLUTIONS LTD.")
    y -= 0.2*inch
    c.drawString(0.75*inch, y, "Thank you for your business!")

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer

# --- Helper: Generate DOCX ---
def build_docx():
    doc = Document()

    # Logo (optional)
    try:
        doc.add_picture("assets/logo.png", width=Inches(1.2))
    except Exception:
        pass

    doc.add_heading("MAKK CROSS BORDER SOLUTIONS LTD.", level=1)
    doc.add_paragraph("address | 626-601-613")

    doc.add_heading("INVOICE", level=2)
    doc.add_paragraph(f"Date: {inv_date.isoformat()}")
    doc.add_paragraph(f"Invoice #: {invoice_no}")

    doc.add_paragraph(f"To: {receiver}")
    doc.add_paragraph(f"Phone: {phone}")
    doc.add_paragraph("Address:")
    doc.add_paragraph(address or "")

    table = doc.add_table(rows=2, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Qty"
    hdr[1].text = "Description"
    hdr[2].text = "Weight"
    hdr[3].text = "Amount"

    row = table.rows[1].cells
    row[0].text = f"{quantity:g}"
    row[1].text = description
    row[2].text = f"{weight:g}"
    row[3].text = f"${line_total:,.2f}"

    doc.add_paragraph(f"Subtotal: ${subtotal:,.2f}")
    doc.add_paragraph(f"Sales Tax: ${sales_tax:,.2f}")
    doc.add_paragraph(f"Total: ${total:,.2f}")

    doc.add_paragraph("Note: Make all checks payable to MAKK CROSS BORDER SOLUTIONS LTD.")
    doc.add_paragraph("Thank you for your business!")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

st.divider()

# Export buttons
pdf_bytes = build_pdf()
st.download_button(
    label="Download PDF",
    data=pdf_bytes,
    file_name=f"MAKK_Invoice_{invoice_no or 'draft'}.pdf",
    mime="application/pdf",
)

docx_bytes = build_docx()
st.download_button(
    label="Download Word (DOCX)",
    data=docx_bytes,
    file_name=f"MAKK_Invoice_{invoice_no or 'draft'}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)
