import io
from datetime import date

import streamlit as st
import pandas as pd

from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle

from docx import Document
from docx.shared import Inches

# ----------------------------
# Page config
# ----------------------------
st.set_page_config(page_title="MAKK Invoice Generator", layout="centered")

# ----------------------------
# Company config (edit here)
# ----------------------------
COMPANY_NAME = "MAKK CROSS BORDER SOLUTIONS LTD."
COMPANY_ADDR = "14278 Valley Blvd.\nLa Puente, CA 91746"
COMPANY_PHONE = "626-601-613"
PAYABLE_NOTE = "Make all checks payable to MAKK CROSS BORDER SOLUTIONS LTD."
THANK_YOU = "Thank you for your business!"

LOGO_PATH = "assets/logo.png"  # put your logo here (optional)


# ----------------------------
# Helpers
# ----------------------------
def money(x: float) -> str:
    return f"${x:,.2f}"


def safe_float(v) -> float:
    try:
        if v is None or v == "":
            return 0.0
        return float(v)
    except Exception:
        return 0.0


def safe_str(v) -> str:
    return "" if v is None else str(v)


# ----------------------------
# UI
# ----------------------------
st.title("MAKK Invoice Generator")

colA, colB = st.columns(2)

with colA:
    inv_date = st.date_input("Date", value=date.today())
    invoice_no = st.text_input("Invoice # (or Customer ID)", value="")

with colB:
    receiver = st.text_input("To (Receiver name / Company)", value="")
    phone = st.text_input("Phone", value="")

address = st.text_area("Address", value="", height=100)

st.subheader("Line Items")

# Use session_state to persist rows
if "items_df" not in st.session_state:
    st.session_state.items_df = pd.DataFrame(
        [
            {"Description": "", "Weight": 0.0, "Amount": 0.0},
        ]
    )

edited_df = st.data_editor(
    st.session_state.items_df,
    use_container_width=True,
    num_rows="dynamic",
    column_config={
        "Description": st.column_config.TextColumn("Description", width="large"),
        "Weight": st.column_config.NumberColumn("Total Weight", min_value=0.0, step=0.1, format="%.2f"),
        "Amount": st.column_config.NumberColumn("Line Total (USD)", min_value=0.0, step=1.0, format="%.2f"),
    },
    hide_index=True,
)

# Clean/normalize
items_df = edited_df.copy()
items_df["Description"] = items_df["Description"].map(safe_str)
items_df["Weight"] = items_df["Weight"].map(safe_float)
items_df["Amount"] = items_df["Amount"].map(safe_float)

# Save back
st.session_state.items_df = items_df

subtotal = float(items_df["Amount"].sum())
sales_tax = st.number_input("Sales Tax (USD) - manual", min_value=0.0, step=1.0, value=0.0)
total = round(subtotal + float(sales_tax), 2)

st.markdown(f"**Subtotal (auto): {money(subtotal)}**")
st.markdown(f"**Total (auto): {money(total)}**")

st.divider()

note_default = f"Note: {PAYABLE_NOTE}\n{THANK_YOU}"
note = st.text_area("Note (shown on invoice)", value=note_default, height=80)


# ----------------------------
# PDF generation (clean invoice layout)
# ----------------------------
def build_pdf() -> io.BytesIO:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)
    w, h = LETTER

    margin_x = 0.75 * inch
    top_y = h - 0.75 * inch

    # --- Logo (optional) ---
    logo_w = 1.1 * inch
    logo_h = 1.1 * inch
    logo_y = top_y - logo_h + 0.15 * inch
    logo_x = margin_x
    logo_drawn = False
    try:
        c.drawImage(LOGO_PATH, logo_x, logo_y, width=logo_w, height=logo_h, mask="auto")
        logo_drawn = True
    except Exception:
        logo_drawn = False

    # --- Company block (CENTER) ---
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(w / 2, top_y, COMPANY_NAME)

    c.setFont("Helvetica", 10)
    # Address + phone centered under company name
    addr_lines = COMPANY_ADDR.split("\n")
    y_company = top_y - 0.2 * inch
    for line in addr_lines:
        c.drawCentredString(w / 2, y_company, line)
        y_company -= 0.18 * inch

    c.drawCentredString(w / 2, y_company, COMPANY_PHONE)
    y_company -= 0.10 * inch

    # --- Right block: INVOICE meta (RIGHT ALIGNED) ---
    # place on the right, aligned
    right_x = w - margin_x
    meta_top = top_y - 0.05 * inch  # slightly below top line

    c.setFont("Helvetica-Bold", 18)
    c.drawRightString(right_x, meta_top - 0.2 * inch, "INVOICE")

    c.setFont("Helvetica-Bold", 10)
    c.drawRightString(right_x, meta_top - 0.50 * inch, "Invoice #")
    c.setFont("Helvetica", 10)
    c.drawRightString(right_x, meta_top - 0.66 * inch, safe_str(invoice_no))

    c.setFont("Helvetica-Bold", 10)
    c.drawRightString(right_x, meta_top - 0.86 * inch, "Date")
    c.setFont("Helvetica", 10)
    c.drawRightString(right_x, meta_top - 1.02 * inch, inv_date.strftime("%Y/%m/%d"))

    # --- Bill to block (LEFT) ---
    y = top_y - 1.55 * inch
    c.setFont("Helvetica-Bold", 11)
    c.drawString(margin_x, y, "BILL TO")
    y -= 0.22 * inch

    c.setFont("Helvetica", 10)
    if receiver.strip():
        c.drawString(margin_x, y, receiver)
        y -= 0.18 * inch
    if phone.strip():
        c.drawString(margin_x, y, phone)
        y -= 0.18 * inch
    if address.strip():
        for line in address.split("\n"):
            c.drawString(margin_x, y, line)
            y -= 0.18 * inch

    # --- Items table ---
    y_table_top = y - 0.25 * inch

    data = [["DESCRIPTION", "WEIGHT", "AMOUNT"]]
    for _, row in items_df.iterrows():
        desc = row["Description"].strip() or "-"
        wt = row["Weight"]
        amt = row["Amount"]
        data.append([desc, f"{wt:,.2f}", money(amt)])

    # table width
    table_w = w - 2 * margin_x
    col_widths = [table_w * 0.64, table_w * 0.18, table_w * 0.18]

    tbl = Table(data, colWidths=col_widths)
    tbl.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 10),
                ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("LINEBELOW", (0, 0), (-1, 0), 1, colors.black),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 1), (-1, -1), 10),
                ("ALIGN", (1, 1), (1, -1), "RIGHT"),
                ("ALIGN", (2, 1), (2, -1), "RIGHT"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )

    # compute table height and draw
    # Use a Platypus frame-like draw on canvas
    table_x = margin_x
    table_y = y_table_top - 0.2 * inch
    _, table_h = tbl.wrapOn(c, table_w, h)
    tbl.drawOn(c, table_x, table_y - table_h)

    # --- Totals block (RIGHT) ---
    y_totals = table_y - table_h - 0.35 * inch
    c.setFont("Helvetica", 10)

    label_x = w - margin_x - 2.2 * inch
    value_x = w - margin_x

    c.drawRightString(label_x, y_totals, "Subtotal")
    c.drawRightString(value_x, y_totals, money(subtotal))
    y_totals -= 0.22 * inch

    c.drawRightString(label_x, y_totals, "Sales Tax")
    c.drawRightString(value_x, y_totals, money(float(sales_tax)))
    y_totals -= 0.22 * inch

    c.setFont("Helvetica-Bold", 11)
    c.drawRightString(label_x, y_totals, "Total")
    c.drawRightString(value_x, y_totals, money(total))
    y_totals -= 0.35 * inch

    # --- Note box (BOTTOM) ---
    # Professional note area
    box_x = margin_x
    box_w = w - 2 * margin_x
    box_h = 0.95 * inch
    box_y = 0.85 * inch

    c.setStrokeColor(colors.black)
    c.rect(box_x, box_y, box_w, box_h, stroke=1, fill=0)

    c.setFont("Helvetica", 10)
    text = c.beginText(box_x + 0.2 * inch, box_y + box_h - 0.28 * inch)
    for line in (note or "").split("\n"):
        text.textLine(line.strip())
    c.drawText(text)

    c.showPage()
    c.save()
    buf.seek(0)
    return buf


# ----------------------------
# DOCX generation (optional)
# ----------------------------
def build_docx() -> io.BytesIO:
    doc = Document()

    # Logo optional
    try:
        doc.add_picture(LOGO_PATH, width=Inches(1.2))
    except Exception:
        pass

    doc.add_paragraph(COMPANY_NAME)
    doc.add_paragraph(COMPANY_ADDR.replace("\n", " "))
    doc.add_paragraph(COMPANY_PHONE)

    doc.add_paragraph("")
    doc.add_paragraph("INVOICE")

    doc.add_paragraph(f"Invoice #: {invoice_no}")
    doc.add_paragraph(f"Date: {inv_date.strftime('%Y/%m/%d')}")

    doc.add_paragraph("")
    doc.add_paragraph("BILL TO")
    doc.add_paragraph(receiver)
    doc.add_paragraph(phone)
    doc.add_paragraph(address)

    doc.add_paragraph("")
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "DESCRIPTION"
    hdr[1].text = "WEIGHT"
    hdr[2].text = "AMOUNT"

    for _, r in items_df.iterrows():
        row = table.add_row().cells
        row[0].text = r["Description"].strip() or "-"
        row[1].text = f"{float(r['Weight']):,.2f}"
        row[2].text = money(float(r["Amount"]))

    doc.add_paragraph("")
    doc.add_paragraph(f"Subtotal: {money(subtotal)}")
    doc.add_paragraph(f"Sales Tax: {money(float(sales_tax))}")
    doc.add_paragraph(f"Total: {money(total)}")

    doc.add_paragraph("")
    doc.add_paragraph(note or "")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ----------------------------
# Download buttons
# ----------------------------
pdf_bytes = build_pdf()
st.download_button(
    "Download PDF",
    data=pdf_bytes,
    file_name=f"MAKK_Invoice_{invoice_no or 'draft'}.pdf",
    mime="application/pdf",
)

docx_bytes = build_docx()
st.download_button(
    "Download Word (DOCX)",
    data=docx_bytes,
    file_name=f"MAKK_Invoice_{invoice_no or 'draft'}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)
